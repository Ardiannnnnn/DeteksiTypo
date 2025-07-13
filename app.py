from flask import Flask, render_template, request, jsonify
import re
from difflib import get_close_matches
import docx
import PyPDF2
import io

app = Flask(__name__)

# Dictionary kata-kata bahasa Indonesia yang umum dan kata-kata akademik
INDONESIAN_WORDS = {
    # Kata umum
    'selamat', 'surat', 'datang', 'halo', 'terima', 'kasih', 'tolong', 'maaf', 'permisi',
    'bagaimana', 'kabar', 'baik', 'buruk', 'cantik', 'indah', 'jelek', 'besar',
    'kecil', 'tinggi', 'rendah', 'panjang', 'pendek', 'lebar', 'sempit', 'tebal',
    'tipis', 'berat', 'ringan', 'keras', 'lunak', 'panas', 'dingin', 'hangat',
    'sejuk', 'segar', 'lezat', 'hambar', 'manis', 'pahit', 'asin', 'asam',
    'pedas', 'gurih', 'wangi', 'bau', 'harum', 'busuk', 'bersih', 'kotor',
    'rapi', 'berantakan', 'terang', 'gelap', 'cerah', 'suram', 'ramai', 'sepi',
    'cepat', 'lambat', 'mudah', 'sulit', 'senang', 'sedih', 'gembira', 'kecewa',
    'marah', 'takut', 'berani', 'malu', 'bangga', 'rendah', 'hati', 'sombong',
    'ramah', 'kasar', 'sopan', 'nakal', 'rajin', 'malas', 'pintar', 'bodoh',
    'pandai', 'dungu', 'cerdas', 'tolol', 'bijak', 'bodoh', 'sabar', 'sabaran',
    
    # Kata akademik dan laporan
    'penelitian', 'analisis', 'metodologi', 'hipotesis', 'kesimpulan', 'pembahasan',
    'pendahuluan', 'abstrak', 'referensi', 'daftar', 'pustaka', 'literatur',
    'tinjauan', 'teori', 'konsep', 'definisi', 'pengertian', 'klasifikasi',
    'kategori', 'pengelompokan', 'identifikasi', 'karakteristik', 'ciri',
    'sifat', 'fungsi', 'manfaat', 'kegunaan', 'tujuan', 'sasaran', 'target',
    'hasil', 'temuan', 'penemuan', 'data', 'informasi', 'fakta', 'bukti',
    'evaluasi', 'penilaian', 'pengukuran', 'observasi', 'pengamatan', 'survei',
    'wawancara', 'kuisioner', 'angket', 'instrumen', 'variabel', 'indikator',
    'parameter', 'kriteria', 'standar', 'norma', 'aturan', 'prosedur',
    'metode', 'teknik', 'cara', 'langkah', 'tahap', 'fase', 'proses',
    'implementasi', 'penerapan', 'aplikasi', 'eksperimen', 'percobaan',
    'simulasi', 'model', 'kerangka', 'struktur', 'sistem', 'skema',
    'diagram', 'grafik', 'tabel', 'chart', 'gambar', 'ilustrasi',
    'perbandingan', 'komparasi', 'perbedaan', 'persamaan', 'hubungan',
    'korelasi', 'relasi', 'kaitan', 'pengaruh', 'dampak', 'efek',
    'faktor', 'elemen', 'komponen', 'bagian', 'aspek', 'dimensi',
    'lingkup', 'ruang', 'waktu', 'temporal', 'spasial', 'geografis',
    'demografi', 'populasi', 'sampel', 'responden', 'subjek', 'objek',
    
    # Kata teknologi
    'program', 'komputer', 'internet', 'website', 'aplikasi', 'software',
    'hardware', 'teknologi', 'digital', 'smartphone', 'laptop', 'desktop',
    'server', 'database', 'python', 'javascript', 'html', 'css', 'flask',
    'framework', 'library', 'function', 'variable', 'algoritma', 'struktur',
    'data', 'array', 'string', 'integer', 'boolean', 'object', 'class',
    'method', 'inheritance', 'polymorphism', 'encapsulation', 'abstraction',
    
    # Kata formal dan akademis tambahan
    'universitas', 'fakultas', 'jurusan', 'program', 'studi', 'mahasiswa',
    'dosen', 'profesor', 'doktor', 'magister', 'sarjana', 'diploma',
    'skripsi', 'tesis', 'disertasi', 'makalah', 'artikel', 'jurnal',
    'seminar', 'konferensi', 'simposium', 'workshop', 'pelatihan',
    'pembelajaran', 'pendidikan', 'pengajaran', 'kurikulum', 'silabus',
    
    # Kata umum tambahan
    'dan', 'atau', 'tetapi', 'namun', 'karena', 'sebab', 'oleh', 'untuk',
    'dari', 'dalam', 'pada', 'di', 'ke', 'dengan', 'tanpa', 'seperti',
    'sama', 'beda', 'lebih', 'kurang', 'sangat', 'sekali', 'juga', 'pula',
    'saja', 'hanya', 'bukan', 'tidak', 'belum', 'sudah', 'akan', 'telah',
    'sedang', 'masih', 'lagi', 'kembali', 'semua', 'setiap', 'masing',
    'sendiri', 'bersama', 'antara', 'selama', 'setelah', 'sebelum', 'ketika',
    'dimana', 'mengapa', 'bagaimana', 'siapa', 'apa', 'kapan', 'berapa',
    
    # Kata kerja umum
    'adalah','apakah', 'ialah', 'yaitu', 'yakni', 'berupa', 'merupakan', 'menjadi',
    'membuat', 'melakukan', 'menggunakan', 'memiliki', 'mempunyai', 'mendapat',
    'memperoleh', 'mencapai', 'menghasilkan', 'memberikan', 'menunjukkan',
    'menjelaskan', 'menyatakan', 'menyebutkan', 'mengatakan', 'berbicara',
    'berkata', 'berpendapat', 'berpikir', 'merasa', 'menganggap', 'melihat',
    'mendengar', 'membaca', 'menulis', 'belajar', 'mengajar', 'bekerja',
    'bermain', 'tidur', 'bangun', 'makan', 'minum', 'pergi', 'datang',
    'pulang', 'masuk', 'keluar', 'naik', 'turun', 'berhenti', 'mulai',
    'selesai', 'berakhir', 'terjadi', 'berlangsung', 'berjalan', 'berlari'
}

def read_docx_file(file_content):
    """Baca file DOCX dengan error handling yang lebih baik"""
    try:
        doc = docx.Document(io.BytesIO(file_content))
        text = []
        
        # Baca paragraf
        for paragraph in doc.paragraphs:
            if paragraph.text.strip():
                text.append(paragraph.text)
        
        # Baca tabel jika ada
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    if cell.text.strip():
                        text.append(cell.text)
        
        result = '\n'.join(text)
        if not result.strip():
            raise Exception("Dokumen DOCX kosong atau tidak berisi teks")
            
        return result
        
    except Exception as e:
        raise Exception(f"Error membaca file DOCX: {str(e)}")

def read_pdf_file(file_content):
    """Baca file PDF dengan mengabaikan gambar dan fokus pada teks"""
    try:
        pdf_reader = PyPDF2.PdfReader(io.BytesIO(file_content))
        text = []
        readable_pages = 0
        
        # Cek apakah PDF memiliki halaman
        if len(pdf_reader.pages) == 0:
            raise Exception("PDF tidak memiliki halaman")
        
        print(f"PDF memiliki {len(pdf_reader.pages)} halaman")
        
        for page_num, page in enumerate(pdf_reader.pages):
            try:
                page_text = page.extract_text()
                
                # Filter dan bersihkan teks
                if page_text and page_text.strip():
                    # Hapus karakter aneh dan non-printable
                    cleaned_text = clean_extracted_text(page_text)
                    if len(cleaned_text.strip()) > 10:  # Minimal 10 karakter
                        text.append(cleaned_text)
                        readable_pages += 1
                        print(f"Halaman {page_num + 1}: berhasil dibaca ({len(cleaned_text)} karakter)")
                    else:
                        print(f"Halaman {page_num + 1}: teks terlalu sedikit, mungkin gambar")
                else:
                    print(f"Halaman {page_num + 1}: tidak ada teks, kemungkinan gambar/scan")
                    
            except Exception as e:
                print(f"Error reading page {page_num + 1}: {str(e)}")
                continue
        
        print(f"Total halaman yang berhasil dibaca: {readable_pages}/{len(pdf_reader.pages)}")
        
        if not text:
            # Jika tidak ada teks sama sekali, coba metode alternatif
            return extract_text_alternative_method(file_content)
        
        result = '\n'.join(text)
        return result
        
    except Exception as e:
        raise Exception(f"Error membaca file PDF: {str(e)}")

def clean_extracted_text(text):
    """Bersihkan teks hasil ekstraksi dari karakter aneh"""
    try:
        # Hapus karakter kontrol dan non-printable
        import string
        printable = set(string.printable)
        
        # Izinkan huruf, angka, spasi, dan tanda baca dasar
        allowed_chars = set(string.ascii_letters + string.digits + string.punctuation + string.whitespace)
        allowed_chars.update('àáâãäåæçèéêëìíîïðñòóôõöøùúûüýþÿ')  # Karakter aksen
        allowed_chars.update('ÀÁÂÃÄÅÆÇÈÉÊËÌÍÎÏÐÑÒÓÔÕÖØÙÚÛÜÝÞŸ')
        
        # Filter karakter
        cleaned = ''.join(char for char in text if char in allowed_chars or char.isalnum())
        
        # Normalize whitespace
        cleaned = re.sub(r'\s+', ' ', cleaned)
        cleaned = re.sub(r'\n+', '\n', cleaned)
        
        return cleaned.strip()
        
    except Exception as e:
        print(f"Error cleaning text: {str(e)}")
        return text

def extract_text_alternative_method(file_content):
    """Metode alternatif untuk ekstraksi teks dengan threshold lebih rendah"""
    try:
        pdf_reader = PyPDF2.PdfReader(io.BytesIO(file_content))
        all_text = []
        
        for page_num, page in enumerate(pdf_reader.pages):
            try:
                # Coba berbagai metode ekstraksi
                page_text = page.extract_text()
                
                # Jika method biasa tidak berhasil, coba extract object text
                if not page_text or len(page_text.strip()) < 5:
                    try:
                        # Ekstrak dari content stream langsung
                        if '/Contents' in page:
                            content = page['/Contents']
                            if hasattr(content, 'get_data'):
                                raw_content = content.get_data().decode('latin-1', errors='ignore')
                                # Cari pola teks sederhana
                                text_matches = re.findall(r'\((.*?)\)', raw_content)
                                if text_matches:
                                    page_text = ' '.join(text_matches)
                    except:
                        pass
                
                if page_text and len(page_text.strip()) > 3:
                    cleaned = clean_extracted_text(page_text)
                    if cleaned:
                        all_text.append(cleaned)
                        print(f"Metode alternatif - Halaman {page_num + 1}: {len(cleaned)} karakter")
                        
            except Exception as e:
                print(f"Metode alternatif gagal pada halaman {page_num + 1}: {str(e)}")
                continue
        
        if all_text:
            return '\n'.join(all_text)
        else:
            # Last resort: return informative message
            raise Exception("PDF berisi mostly gambar/scan. Silakan convert ke teks atau gunakan file format lain")
            
    except Exception as e:
        raise Exception(f"Metode alternatif gagal: {str(e)}")

def detect_typos(text):
    """Deteksi typo dalam teks"""
    # Bersihkan teks dan ambil kata-kata
    words = re.findall(r'\b[a-zA-ZÀ-ÿ]+\b', text.lower())
    typos = []
    
    for word in words:
        # Skip kata yang sangat pendek (1-2 huruf)
        if len(word) <= 2:
            continue
            
        if word not in INDONESIAN_WORDS:
            # Cari kata yang mirip dengan threshold yang lebih ketat
            suggestions = get_close_matches(word, INDONESIAN_WORDS, n=3, cutoff=0.7)
            
            # Hanya tambahkan jika ada saran atau kata cukup panjang untuk dicurigai typo
            if suggestions or len(word) >= 4:
                typos.append({
                    'word': word,
                    'suggestions': suggestions
                })
    
    return typos

def read_txt_file(file_content):
    """Baca file TXT dengan berbagai encoding"""
    try:
        # Coba beberapa encoding
        encodings = ['utf-8', 'utf-8-sig', 'latin-1', 'cp1252', 'iso-8859-1']
        
        for encoding in encodings:
            try:
                text = file_content.decode(encoding)
                if text.strip():  # Pastikan ada isi
                    return text
            except UnicodeDecodeError:
                continue
        
        raise Exception("Tidak dapat membaca file dengan encoding yang didukung")
        
    except Exception as e:
        raise Exception(f"Error membaca file TXT: {str(e)}")

def is_text_meaningful(text):
    """Periksa apakah teks yang diekstrak bermakna (bukan noise)"""
    if not text or len(text.strip()) < 10:
        return False
    
    # Hitung rasio huruf vs karakter aneh
    letter_count = sum(1 for c in text if c.isalpha())
    total_count = len(text)
    
    if total_count == 0:
        return False
    
    letter_ratio = letter_count / total_count
    
    # Jika kurang dari 30% huruf, kemungkinan noise
    if letter_ratio < 0.3:
        return False
    
    # Periksa apakah ada kata-kata yang masuk akal
    words = re.findall(r'\b[a-zA-Z]{3,}\b', text)
    if len(words) < 3:  # Minimal 3 kata
        return False
    
    return True

@app.route('/')
def home():
    """Halaman utama"""
    return render_template('index.html')

@app.route('/check_typo', methods=['POST'])
def check_typo():
    """API endpoint untuk memeriksa typo"""
    try:
        data = request.get_json()
        text = data.get('text', '')
        
        if not text:
            return jsonify({'error': 'Teks tidak boleh kosong'}), 400
        
        if len(text.strip()) < 3:
            return jsonify({'error': 'Teks terlalu pendek untuk diperiksa'}), 400
        
        typos = detect_typos(text)
        
        return jsonify({
            'typos': typos,
            'total_typos': len(typos),
            'message': f"Ditemukan {len(typos)} potensi typo dalam teks Anda" if typos else "Tidak ada typo yang terdeteksi!"
        })
        
    except Exception as e:
        return jsonify({'error': f'Terjadi kesalahan: {str(e)}'}), 500

@app.route('/upload_file', methods=['POST'])
def upload_file():
    """API endpoint untuk upload dan membaca file"""
    try:
        if 'file' not in request.files:
            return jsonify({'error': 'Tidak ada file yang diupload'}), 400
        
        file = request.files['file']
        if file.filename == '':
            return jsonify({'error': 'Tidak ada file yang dipilih'}), 400
        
        # Baca konten file
        file_content = file.read()
        filename = file.filename.lower()
        
        # Debug: Print info file
        print(f"Filename: {filename}")
        print(f"File size: {len(file_content)} bytes")
        
        text = ""
        file_type = ""
        
        # Tentukan jenis file dan baca sesuai format
        if filename.endswith('.txt'):
            file_type = "TXT"
            text = read_txt_file(file_content)
        elif filename.endswith('.docx'):
            file_type = "DOCX"
            text = read_docx_file(file_content)
        elif filename.endswith('.pdf'):
            file_type = "PDF"
            try:
                text = read_pdf_file(file_content)
            except Exception as pdf_error:
                # Jika PDF berisi scan/gambar, berikan pesan yang jelas
                if "mostly gambar" in str(pdf_error) or "scan" in str(pdf_error):
                    return jsonify({
                        'error': 'PDF berisi gambar/scan dan tidak dapat dibaca secara otomatis.',
                        'suggestion': 'Silakan copy-paste teks secara manual ke input teks atau convert PDF ke format Word/TXT terlebih dahulu.'
                    }), 400
                else:
                    raise pdf_error
        elif filename.endswith('.doc'):
            return jsonify({'error': 'File DOC lama tidak didukung. Gunakan DOCX, TXT, atau PDF'}), 400
        else:
            return jsonify({'error': 'Format file tidak didukung. Gunakan TXT, DOCX, atau PDF'}), 400
        
        # Debug: Print extracted text info
        print(f"Extracted text length: {len(text) if text else 0}")
        print(f"First 200 chars: {text[:200] if text else 'No text'}")
        
        # Validasi teks yang diekstrak
        if not text or not text.strip():
            return jsonify({
                'error': f'File {file_type} kosong atau tidak berisi teks yang dapat dibaca.'
            }), 400
        
        if len(text.strip()) < 20:
            return jsonify({
                'error': f'Teks terlalu pendek untuk dianalisis ({len(text.strip())} karakter). Minimal 20 karakter diperlukan.'
            }), 400
        
        # Validasi meaningful text hanya untuk PDF (karena bisa berisi noise)
        if file_type == "PDF" and not is_text_meaningful(text):
            return jsonify({
                'error': f'PDF tidak berisi teks yang bermakna. Kemungkinan file berisi gambar/scan.',
                'suggestion': 'Silakan copy-paste teks secara manual atau gunakan file yang berisi teks asli.'
            }), 400
        
        # Deteksi typo
        typos = detect_typos(text)
        
        return jsonify({
            'text': text,
            'typos': typos,
            'total_typos': len(typos),
            'file_type': file_type,
            'text_length': len(text),
            'message': f"File {file_type} berhasil diproses. Ditemukan {len(typos)} potensi typo." if typos else f"File {file_type} berhasil diproses. Tidak ada typo yang terdeteksi!"
        })
        
    except Exception as e:
        print(f"Error in upload_file: {str(e)}")
        return jsonify({'error': f'Terjadi kesalahan: {str(e)}'}), 500

@app.errorhandler(404)
def not_found(error):
    """Handler untuk halaman tidak ditemukan"""
    return render_template('index.html'), 404

@app.errorhandler(500)
def internal_error(error):
    """Handler untuk error server"""
    return jsonify({'error': 'Terjadi kesalahan internal server'}), 500


if __name__ == '__main__':
    # Development mode dengan auto-reload yang lebih proper
    app.run(
        debug=True,              # Enable debug mode
        host='127.0.0.1',        # Localhost untuk development
        port=5000,               # Port number
        use_reloader=True,       # Auto-reload on file changes
        use_debugger=True,       # Enable debugger
        threaded=True            # Enable threading
    )